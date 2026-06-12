"""Convert SETUP.md into a formatted Word .docx (headings, tables, code, lists)."""
import re
import sys
from docx import Document
from docx.shared import Pt, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml.ns import qn
from docx.oxml import OxmlElement

SRC = sys.argv[1] if len(sys.argv) > 1 else "SETUP.md"
OUT = sys.argv[2] if len(sys.argv) > 2 else "PoultryMaster Setup Guide.docx"


def shade(paragraph, fill="F2F2F2"):
    p = paragraph._p
    pPr = p.get_or_add_pPr()
    shd = OxmlElement("w:shd")
    shd.set(qn("w:val"), "clear")
    shd.set(qn("w:fill"), fill)
    pPr.append(shd)


def add_runs(paragraph, text):
    """Render inline **bold** and `code` into a paragraph."""
    # split on bold or inline code, keeping delimiters
    parts = re.split(r"(\*\*[^*]+\*\*|`[^`]+`)", text)
    for part in parts:
        if not part:
            continue
        if part.startswith("**") and part.endswith("**"):
            r = paragraph.add_run(part[2:-2]); r.bold = True
        elif part.startswith("`") and part.endswith("`"):
            r = paragraph.add_run(part[1:-1])
            r.font.name = "Consolas"; r.font.size = Pt(9.5)
            r.font.color.rgb = RGBColor(0xC0, 0x39, 0x2B)
        else:
            paragraph.add_run(part)


def add_code_block(doc, lines):
    for ln in lines:
        p = doc.add_paragraph()
        p.paragraph_format.left_indent = Pt(8)
        p.paragraph_format.space_after = Pt(0)
        p.paragraph_format.space_before = Pt(0)
        shade(p)
        r = p.add_run(ln if ln else " ")
        r.font.name = "Consolas"; r.font.size = Pt(9)


def add_table(doc, rows):
    cells = [[c.strip() for c in r.strip().strip("|").split("|")] for r in rows]
    cells = [c for i, c in enumerate(cells) if not re.match(r"^[-:\s|]+$", "|".join(c))]
    if not cells:
        return
    ncol = max(len(r) for r in cells)
    t = doc.add_table(rows=len(cells), cols=ncol)
    t.style = "Light Grid Accent 1"
    for ri, row in enumerate(cells):
        for ci in range(ncol):
            cell = t.cell(ri, ci)
            cell.text = ""
            p = cell.paragraphs[0]
            add_runs(p, row[ci] if ci < len(row) else "")
            if ri == 0:
                for run in p.runs:
                    run.bold = True
    doc.add_paragraph()


def main():
    with open(SRC, encoding="utf-8") as f:
        lines = f.read().splitlines()

    doc = Document()
    style = doc.styles["Normal"]
    style.font.name = "Calibri"; style.font.size = Pt(11)

    i = 0
    n = len(lines)
    while i < n:
        line = lines[i]

        # fenced code block
        if line.strip().startswith("```"):
            block = []
            i += 1
            while i < n and not lines[i].strip().startswith("```"):
                block.append(lines[i]); i += 1
            add_code_block(doc, block)
            i += 1
            continue

        # table
        if line.strip().startswith("|"):
            block = []
            while i < n and lines[i].strip().startswith("|"):
                block.append(lines[i]); i += 1
            add_table(doc, block)
            continue

        stripped = line.strip()

        if stripped.startswith("# "):
            h = doc.add_heading(level=0); add_runs(h, stripped[2:])
        elif stripped.startswith("## "):
            h = doc.add_heading(level=1); add_runs(h, stripped[3:])
        elif stripped.startswith("### "):
            h = doc.add_heading(level=2); add_runs(h, stripped[4:])
        elif stripped == "---":
            pass  # skip horizontal rules
        elif stripped.startswith("> "):
            p = doc.add_paragraph(style="Intense Quote"); add_runs(p, stripped[2:])
        elif re.match(r"^[-*] ", stripped):
            p = doc.add_paragraph(style="List Bullet"); add_runs(p, stripped[2:])
        elif re.match(r"^\d+\. ", stripped):
            p = doc.add_paragraph(style="List Number"); add_runs(p, re.sub(r"^\d+\.\s*", "", stripped))
        elif stripped == "":
            pass
        else:
            p = doc.add_paragraph(); add_runs(p, line)
        i += 1

    doc.save(OUT)
    print("WROTE", OUT)


if __name__ == "__main__":
    main()
